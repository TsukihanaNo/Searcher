#ver 0.27
import os, sys, re
from PySide import QtGui, QtCore
import time
from InfoWindow import *
from DiskReportWindow import *
from SettingWindow import *
from ProcessThread import *
from MonitorGraph import *
from UpdateThread import *
from AppTable import *
from TextEditor import *
from ImageView import *
import docx
#from MusicPlayer import *
if getattr(sys, 'frozen', False):
    # frozen
    program_location = os.path.dirname(sys.executable)
else:
    # unfrozen
    program_location = os.path.dirname(os.path.realpath(__file__))

archivefolder = os.path.join(program_location, 'Archive')
if not os.path.isdir(archivefolder):
    os.mkdir(archivefolder)

file_count = os.path.join(program_location, archivefolder, 'counts.txt')
export_file = os.path.join(program_location, archivefolder, 'foundfiles.txt')
search_locations = os.path.join(program_location, archivefolder, 'locations.txt')
path_archive = os.path.join(program_location, archivefolder, 'path archive.txt')
restricted_file = os.path.join(program_location, archivefolder, 'restricted folders.txt')
error_log_file = os.path.join(program_location, archivefolder, 'error log.txt')
exception_file = os.path.join(program_location, archivefolder, 'exception.txt')
monitor_archive = os.path.join(program_location, 'monitor archive')
timer_setting = os.path.join(program_location,archivefolder,'timer.txt')
#word_list_file = os.path.join(program_location,archivefolder,'word list.txt')

if not os.path.isdir(monitor_archive):
    os.mkdir(monitor_archive)

icon = os.path.join(program_location, 'find.png')
help_file = os.path.join(program_location, 'help.txt')

MAXFILELOAD = 1024 * 15000
BLOCKSIZE = 1024 * 5000
MUSIC_EXT = ['.mp3','.flac','.wav','.wma','.m4a','.aiff','.m4p']
IMAGE_EXT = ['.jpeg','.jpg','.bmp','.tiff','.png','.psd','gif','.jfif','.exif','.']
VIDEO_EXT = ['.avi','.mp4','.mpeg','.mov','.wma','.wmv','.wmx','.ogm','.mkv']
EXT_LIST = ['','.txt', '.ini', '.py', '.htm,', '.html', '.thtml', '.php', '.sql','.c','.cpp','.java', '.css']



class Searcher(QtGui.QWidget):
    def __init__(self):
        super(Searcher, self).__init__()
        # variables
        self.is_log_open = False
        self.is_setting_open = False
        self.is_diskreport_open = False
        self.is_monitor_open = False
        self.is_preview_open = False
        #self.is_music_open = False
        self.restricted_dir_count = 0

        # timer
        self.timer = QtCore.QTimer(self)

        # lists
        self.error_list = []
        self.restricted_dir = []
        self.found = []
        self.root_search_list = []
        self.search_exception_list = []
        self.pathlist = []


        # creating widgets
        self.file_lcd = QtGui.QLCDNumber(self)
        self.dir_lcd = QtGui.QLCDNumber(self)
        self.progress_bar = QtGui.QProgressBar(self)
        self.menu_bar = QtGui.QMenuBar(self)
        self.status_label = QtGui.QLabel('Status: ', self)
        self.line_edit = QtGui.QLineEdit(self)
        self.deep_line = QtGui.QLineEdit(self)
        self.listwidget = QtGui.QListWidget(self)
        self.generate = QtGui.QPushButton('Generate', self)
        self.index_date = QtGui.QLabel(self)

        self.initUI()


    def initUI(self):
        ico = QtGui.QIcon(icon)

        # menu creations
        self.CreateActionsAndMenu()

        # layouts
        vlaybox = QtGui.QVBoxLayout(self)
        hlaybox = QtGui.QHBoxLayout(self)

        v1 = QtGui.QVBoxLayout(self)
        group_box1 = QtGui.QGroupBox('Search', self)
        v1.addWidget(self.line_edit)
        group_box1.setLayout(v1)

        v2 = QtGui.QVBoxLayout(self)
        group_box2 = QtGui.QGroupBox('In File Search', self)
        v2.addWidget(self.deep_line)
        group_box2.setLayout(v2)

        h = QtGui.QHBoxLayout(self)
        group_box3 = QtGui.QGroupBox('Misc', self)
        h.addWidget(self.generate)
        group_box3.setLayout(h)

        h1 = QtGui.QHBoxLayout(self)
        group_box4 = QtGui.QGroupBox('Folders', self)
        h1.addWidget(self.dir_lcd)
        group_box4.setLayout(h1)

        h2 = QtGui.QHBoxLayout(self)
        group_box5 = QtGui.QGroupBox('Files', self)
        h2.addWidget(self.file_lcd)
        group_box5.setLayout(h2)

        hlaybox.addWidget(group_box1)
        hlaybox.addWidget(group_box2)
        hlaybox.addWidget(group_box4)
        hlaybox.addWidget(group_box5)
        hlaybox.addWidget(group_box3)

        vlaybox.setMenuBar(self.menu_bar)
        vlaybox.addLayout(hlaybox)
        vlaybox.addWidget(self.status_label)
        vlaybox.addWidget(self.progress_bar)
        vlaybox.addWidget(self.listwidget)
        vlaybox.addWidget(self.index_date)

        #widget settings
        self.file_lcd.display(0)
        self.dir_lcd.display(0)
        self.dir_lcd.setNumDigits(7)
        self.file_lcd.setNumDigits(9)
        self.status_label.setFixedWidth(550)

        self.progress_bar.setTextVisible(False)
        self.progress_bar.setFixedHeight(3)

        self.dir_lcd.setSegmentStyle(self.dir_lcd.Flat)
        self.file_lcd.setSegmentStyle(self.file_lcd.Flat)
        self.listwidget.setSelectionMode(QtGui.QAbstractItemView.ExtendedSelection)

        #connecting signals
        self.line_edit.returnPressed.connect(self.SearchSplit)
        self.listwidget.doubleClicked.connect(self.OpenFile)
        self.generate.clicked.connect(self.GenerateArchive)
        self.deep_line.returnPressed.connect(self.DeepSearch)
        self.timer.timeout.connect(self.Update)
        self.listwidget.itemSelectionChanged.connect(self.previewing)

        shortcut = QtGui.QShortcut(QtGui.QKeySequence(QtCore.Qt.Key_Backspace),self.listwidget)
        shortcut.activated.connect(self.delete)

        shortcut2 = QtGui.QShortcut(QtGui.QKeySequence(QtCore.Qt.Key_Delete),self.listwidget)
        shortcut2.activated.connect(self.delete)


        if os.path.isfile(file_count):
            f = open(file_count, 'r')
            x = 0
            for line in f:
                if x == 0:
                    self.dir_count = int(line.rstrip('\n'))
                    self.dir_lcd.display(self.dir_count)
                else:
                    self.file_count = int(line.rstrip('\n'))
                    self.file_lcd.display(self.file_count)
                x += 1
            f.close()

        #window settings
        self.setGeometry(100, 50, 675, 100)
        self.setWindowTitle("Searcher")
        self.setAttribute(QtCore.Qt.WA_DeleteOnClose)
        self.setWindowIcon(ico)
        self.show()
        self.Initialize()


    def Initialize(self):
        if os.path.isfile(path_archive):
            self.ReadArchive()
            self.index_date.setText('Index -> Last updated: ' + self.pathlist[len(self.pathlist)-1])
        else:
            self.status_label.setText('Status: There is no stored Index, please generate')
        if not os.path.isfile(exception_file):
            f = open(exception_file, 'w+')
            f.write('')
            f.close()
        if os.path.isfile(timer_setting):
            self.readTimer()

    def delete(self):
        del_list = []
        for item in self.listwidget.selectedItems():
            del_list.append(item.text())
        more_info = ""
        for path in del_list:
            more_info += path + "\n"
        if del_list != []:
            messagebox = QtGui.QMessageBox()
            messagebox.setText("Are you sure you want to delete these items?")
            messagebox.setDetailedText("** "+more_info)
            messagebox.setStandardButtons(QtGui.QMessageBox.Yes|QtGui.QMessageBox.No)
            ret = messagebox.exec_()
            if ret == QtGui.QMessageBox.Yes:
                for path in del_list:
                    if os.path.isfile(path):
                        os.remove(path)
                        self.pathlist.remove(path)
                        self.found.remove(path)
                        self.file_count -=1
                    if os.path.isdir(path):
                        if os.listdir(path)==[]:
                            os.rmdir(path)
                            self.pathlist.remove(path)
                            self.found.remove(path)
                            self.dir_count -=1
                        else:
                            self.deleteDir(path)
                            os.rmdir(path)
                            self.pathlist.remove(path)
                            self.found.remove(path)
                            self.dir_count -=1
                self.listwidget.clear()
                for item in self.found:
                    self.listwidget.addItem(item)
                self.status_label.setText("Status: Found " + str(len(self.found)) + " items")
                self.file_lcd.display(self.file_count)
                self.dir_lcd.display(self.dir_count)
                self.CreateArchive()


    def deleteDir(self,path):
        for item in os.listdir(path):
            item = os.path.join(path,item)
            if os.path.isfile(item):
                os.remove(item)
                self.pathlist.remove(item)
                if item in self.found:
                    self.found.remove(item)
                self.file_count -=1
            else:
                if os.listdir(item)==[]:
                    os.rmdir(item)
                    self.pathlist.remove(item)
                    if item in self.found:
                        self.found.remove(item)
                    self.dir_count-=1
                else:
                    self.deleteDir(item)


    def readTimer(self):
        f = open(timer_setting,'r')
        time = []
        for line in f:
            time.append(line.rstrip('\n'))
        print(time[0],time[1])
        if int(time[0])>0:
            self.timer.setInterval(int(time[0]))
            self.timer.start()
        if time[1] =='on':
            self.Update()
        f.close()

    def setTimer(self):
        f = open(timer_setting,'r')
        time = []
        for line in f:
            time.append(line.rstrip('\n'))
        if int(time[0])>0:
            self.timer.setInterval(int(time[0]))
            print(time[0])
            self.timer.start()
        else:
            self.timer.stop()
        f.close()

    def closeEvent(self, event):
        if self.is_setting_open:
            self.setting.close()
        if self.is_log_open:
            self.info.close()
        if self.is_diskreport_open:
            self.disk_report.close()
        if self.is_monitor_open:
            self.monitor.close()


    def CreateActionsAndMenu(self):
        exitaction = QtGui.QAction('Exit', self)
        exitaction.setShortcut('Ctrl+Q')
        exitaction.triggered.connect(self.close)

        ExportAction = QtGui.QAction('Export found files', self)
        ExportAction.setShortcut('Ctrl+S')
        ExportAction.triggered.connect(self.ExportFound)

        ShowExportAction = QtGui.QAction('Open Recent Exported File', self)
        ShowExportAction.setShortcut('Ctrl+O')
        ShowExportAction.triggered.connect(self.ShowExport)

        ShowLogAction = QtGui.QAction('Open Log', self)
        ShowLogAction.setShortcut('Ctrl+L')
        ShowLogAction.triggered.connect(self.LaunchInfo)

        ShowRestrictedAction = QtGui.QAction('Show Inaccessible Folders', self)
        ShowRestrictedAction.triggered.connect(self.OpenRestricted)

        ResetAction = QtGui.QAction('Reset', self)
        ResetAction.setShortcut('Ctrl+R')
        ResetAction.triggered.connect(self.Reset)

        CopyAction = QtGui.QAction('Copy file to target location', self)
        CopyAction.triggered.connect(self.CopyFiles)

        MoveAction = QtGui.QAction('Move file to target location', self)
        MoveAction.triggered.connect(self.MoveFiles)

        ShowDirAction = QtGui.QAction('Show folders only', self)
        ShowDirAction.triggered.connect(self.ShowDirOnly)

        ShowFilesAction = QtGui.QAction('Show files only', self)
        ShowFilesAction.triggered.connect(self.ShowFilesOnly)

        ShowAllAction = QtGui.QAction('Show all', self)
        ShowAllAction.triggered.connect(self.ShowAll)

        OpenFileAction = QtGui.QAction('Open File', self)
        OpenFileAction.triggered.connect(self.OpenFile)

        OpenContainFolderAction = QtGui.QAction('Open containing folder', self)
        OpenContainFolderAction.triggered.connect(self.OpenContainerFolder)

        # ShowHelpAction = QtGui.QAction('Manual', self)
        # self.ShowHelpAction.triggered.connect(self.ShowManual)

        # ShowChangeAction = QtGui.QAction('Change Log', self)
        # self.ShowChangeAction.triggered.connect(self.ShowChangeLog)

        ShowDiskReportAction = QtGui.QAction('Disk Report', self)
        ShowDiskReportAction.setShortcut("Ctrl+D")
        ShowDiskReportAction.triggered.connect(self.LaunchDiskReport)

        LaunchSettingAction = QtGui.QAction('Settings', self)
        LaunchSettingAction.setShortcut("Ctrl+E")
        LaunchSettingAction.triggered.connect(self.LaunchSettings)

        LaunchMonitorAction = QtGui.QAction('Monitor', self)
        LaunchMonitorAction.setShortcut("Ctrl+M")
        LaunchMonitorAction.triggered.connect(self.LaunchMonitor)

        LaunchAppTableAction = QtGui.QAction('Desktop Table', self)
        LaunchAppTableAction.setShortcut("Ctrl+T")
        LaunchAppTableAction.triggered.connect(self.LaunchAppTable)

        LaunchTextAction = QtGui.QAction('Text Editor', self)
        LaunchTextAction.setShortcut('Ctrl+G')
        LaunchTextAction.triggered.connect(self.LaunchTextEditor)

        grepAction = QtGui.QAction('Grep',self)
        grepAction.triggered.connect(self.grep)

        openHighlightAction = QtGui.QAction('Open and Highlight',self)
        openHighlightAction.triggered.connect(self.openText)

        previewAction = QtGui.QAction("Preview",self)
        previewAction.triggered.connect(self.preview)

        deleteAction = QtGui.QAction("&Delete",self)
        deleteAction.triggered.connect(self.delete)

        #playAction = QtGui.QAction('play',self)
        #playAction.triggered.connect(self.playMusic)

        # test = QtGui.QAction('test',self)
        # test.triggered.connect(self.createWordList)


        # menu bar
        filemenu = self.menu_bar.addMenu('&File')
        viewmenu = self.menu_bar.addMenu('&View')
        modulemenu = self.menu_bar.addMenu('&Modules')
        #helpmenu = self.menu_bar.addMenu('&Help')

        filemenu.addAction(ExportAction)
        filemenu.addAction(ShowExportAction)
        filemenu.addAction(ShowRestrictedAction)
        filemenu.addSeparator()
        filemenu.addAction(LaunchSettingAction)
        filemenu.addSeparator()
        filemenu.addAction(ResetAction)
        filemenu.addSeparator()
        filemenu.addAction(exitaction)

        viewmenu.addAction(ShowLogAction)
        viewmenu.addAction(LaunchAppTableAction)
        viewmenu.addAction(LaunchTextAction)

        modulemenu.addAction(LaunchMonitorAction)
        modulemenu.addAction(ShowDiskReportAction)

        #helpmenu.addAction(ShowHelpAction)
        #helpmenu.addAction(ShowChangeAction)

        # listwidget context menu
        separator = QtGui.QAction(self)
        separator.setSeparator(True)
        separator2 = QtGui.QAction(self)
        separator2.setSeparator(True)

        self.listwidget.setContextMenuPolicy(QtCore.Qt.ActionsContextMenu)
        self.listwidget.addAction(previewAction)
        #self.listwidget.addAction(playAction)
        self.listwidget.addAction(OpenFileAction)
        self.listwidget.addAction(OpenContainFolderAction)
        self.listwidget.addAction(openHighlightAction)
        self.listwidget.addAction(grepAction)
        self.listwidget.addAction(separator)
        self.listwidget.addAction(ShowAllAction)
        self.listwidget.addAction(ShowDirAction)
        self.listwidget.addAction(ShowFilesAction)
        self.listwidget.addAction(separator2)
        self.listwidget.addAction(CopyAction)
        self.listwidget.addAction(MoveAction)
        self.listwidget.addAction(deleteAction)
        # self.listwidget.addAction(test)

    def previewing(self):
        if self.is_preview_open:
            for item in self.listwidget.selectedItems():
                image = item.text()
            if os.path.splitext(image)[-1].lower() in IMAGE_EXT:
                self.imageview.open(image)
        # if self.is_music_open:
        #     for item in self.listwidget.selectedItems():
        #         music = item.text()
        #     if os.path.splitext(music)[-1].lower() in MUSIC_EXT:
        #         self.musicplayer.startMusic(music)

    def preview(self):
        if not self.is_preview_open:
            self.is_preview_open = True
            self.imageview = ImageView()
            self.imageview.destroyed.connect(self.previewClosed)
            for item in self.listwidget.selectedItems():
                image = item.text()
            self.imageview.open(image)
            self.imageview.show()

    def previewClosed(self):
        self.is_preview_open = False

    # def playMusic(self):
    #     if not self.is_music_open:
    #         self.is_music_open = True
    #         self.musicplayer = MusicPlayer()
    #         self.musicplayer.destroyed.connect(self.playerClosed)
    #         for item in self.listwidget.selectedItems():
    #             music = item.text()
    #             print(music)
    #         self.musicplayer.startMusic(music)
    #     else:
    #         self.is_music_open = False

    # def playerClosed(self):
    #     self.is_music_open = False


    def LaunchTextEditor(self):
        self.texteditor = TextEditor()
        self.texteditor.show()

    def LaunchAppTable(self):
        self.table = AppTable()
        self.table.show()

    def LaunchMonitor(self):
        if not self.is_monitor_open:
            self.is_monitor_open = True
            self.monitor = MonitorGraph()
            self.monitor.destroyed.connect(self.CloseMonitor)
            self.monitor.show()
        else:
            self.is_monitor_open = False

    def CloseMonitor(self):
        self.is_monitor_open = False


    def Update(self):
        self.status_label.setText("Status: Updating Index")
        thread = UpdateThread()
        self.pathlist, self.file_count, self.dir_count = thread.run()
        self.file_lcd.display(self.file_count)
        self.dir_lcd.display(self.dir_count)
        self.status_label.setText('Status: Index Updated!')
        self.index_date.setText('Index -> Last Updated: ' + self.pathlist[len(self.pathlist)-1])

    def TotalSize(self):
        size = 0
        for path in self.pathlist:
            if not os.path.isdir(path):
                size += os.path.getsize(path)
        return size

    def LaunchSettings(self):
        pos = self.pos()
        if not self.is_setting_open:
            self.is_setting_open = True
            self.setting = SettingsWindow(pos.x() + 8, pos.y() + self.geometry().height())
            self.setting.destroyed.connect(self.CloseSetting)
            self.setting.time_button.clicked.connect(self.setTimer)
        else:
            self.is_setting_open = False
            self.setting.close()

    def CloseSetting(self):
        self.setTimer()
        self.is_setting_open = False

    def LaunchDiskReport(self):
        if not self.is_diskreport_open:
            self.is_diskreport_open = True
            self.disk_report = DiskReport(self.pos().x(), self.pos().y())
            self.disk_report.destroyed.connect(self.DiskReportClosed)
        else:
            self.is_diskreport_open = False
            self.disk_report.close()
        self.disk_report.show()

    def DiskReportClosed(self):
        self.is_diskreport_open = False


    def CopyFiles(self):
        dialog = QtGui.QFileDialog(self)
        target_dir = dialog.getExistingDirectory()
        processed_data = 0
        total_data = 0
        for item in self.listwidget.selectedItems():
            total_data += os.path.getsize(item.text())
        for item in self.listwidget.selectedItems():
            if sys.platform == 'win32':
                filename = item.text().split('\\')[-1]
            else:
                filename = item.text().split('/')[1]

            path_to = os.path.join(target_dir, filename)
            file_from = open(item.text(), 'rb')
            bytes_to = open(path_to, 'wb')
            if os.path.getsize(item.text()) <= MAXFILELOAD:
                try:
                    bytes_from = file_from.read()
                    bytes_to.write(bytes_from)
                except:
                    if self.is_log_open:
                        self.info.textedit.append('error copying file:' + item.text() + '--skipped')
                        self.info.textedit.append(sys.exc_info()[0], sys.exc_info()[1])
                    self.error_list.append('error copying file:' + item.text())
            else:
                while True:
                    try:
                        bytes_from = file_from.read(BLOCKSIZE)
                        if not bytes_from: break
                        bytes_to.write(bytes_from)
                    except:
                        if self.is_log_open:
                            self.info.textedit.append('error copying file:' + item.text() + '--skipped')
                            self.info.textedit.append(sys.exc_info()[0], sys.exc_info()[1])
                        self.error_list.append('error copying file:' + item.text())
            processed_data += os.path.getsize(item.text())
            self.status_label.setText(
                'Status: Copying Progress: ' + str(int((processed_data / total_data) * 100)) + '%')
            QtGui.QApplication.processEvents()
            self.progress_bar.setValue(int(processed_data / total_data) * 100)

            file_from.close()
            bytes_to.close()

    def MoveFiles(self):
        dialog = QtGui.QFileDialog(self)
        target_dir = dialog.getExistingDirectory()
        counter = 0
        for item in self.listwidget.selectedItems():
            if sys.platform == 'win32':
                filename = item.text().split('\\')[-1]
            else:
                filename = item.text().split('/')[1]
            path_to = os.path.join(target_dir, filename)
            try:
                os.renames(item.text(), path_to)
                self.pathlist.remove(item.text())
                self.pathlist.append(path_to)
            except:
                if self.is_log_open:
                    self.info.textedit.append('error moving file:' + item.text() + '--skipped')
                    self.info.textedit.append(sys.exc_info()[0], sys.exc_info()[1])
                self.error_list.append('error moving file:' + item.text())
            counter += 1
            self.status_label.setText(
                'Status: Moving Progess -> ' + str(counter) + ' out of ' + str(len(self.listwidget.selectedItems())))
            QtGui.QApplication.processEvents()
            self.progress_bar.setValue(100 * (counter / len(self.listwidget.selectedItems())))


    def OpenContainerFolder(self):
        for item in self.listwidget.selectedItems():
            folder = ''
            if sys.platform == 'win32':
                path = item.text().split('\\')
            else:
                path = item.text().split('/')
            for x in range(len(path) - 1):
                folder = os.path.join(folder, path[x])
            if sys.platform == 'win32':
                os.startfile(folder)
            else:
                subprocess.call(['xdg-open', '/' + folder])


    def ShowDirOnly(self):
        self.listwidget.clear()
        for path in self.found:
            if not os.path.isfile(path):
                self.listwidget.addItem(path)


    def ShowFilesOnly(self):
        self.listwidget.clear()
        for path in self.found:
            if os.path.isfile(path):
                self.listwidget.addItem(path)


    def ShowAll(self):
        self.listwidget.clear()
        for path in self.found:
            QtGui.QApplication.processEvents()
            self.listwidget.addItem(path)

    def Reset(self):
        f = open(locations, 'w+')
        f.write('')
        f.close()
        f = open(exception_file, 'w+')
        f.write('')
        f.close()
        f = open(file_count, 'w+')
        f.write('')
        f.close()
        self.dir_lcd.display(0)
        self.file_lcd.display(0)

    def LaunchInfo(self):
        self.is_log_open = True
        pos = self.pos()
        if sys.platform == 'win32':
            self.info = InfoWindow(pos.x() + self.geometry().width() + 25, pos.y() + 30)
        else:
            self.info = InfoWindow(pos.x() + self.geometry().width() + 25, pos.y())
        self.info.destroyed.connect(self.InfoClosed)
        self.info.show()
        self.info.textedit.clear()
        self.info.textedit.append('Errors: ' + str(len(self.error_list)) + '\n')
        for item in self.error_list:
            self.info.textedit.append(item)
        self.info.textedit.append('\n')
        self.info.textedit.append('Folders: Could not Access:  ' + str(self.restricted_dir_count) + '\n')
        for item in self.restricted_dir:
            self.info.textedit.append(item)


    def InfoClosed(self):
        self.is_log_open = False


    def ExportFound(self):
        file_dialog = QtGui.QFileDialog(self)
        target = file_dialog.getSaveFileName()
        f = open(target[0], 'w+')
        try:
            for item in self.found:
                f.write(item)
                f.write('\n')
        except UnicodeEncodeError:
            print('error writing')
        f.close()
        f = open(export_file, 'w+')
        try:
            for item in self.found:
                f.write(item)
                f.write('\n')
        except UnicodeEncodeError:
            print('error writing')
        f.close()

    def ShowExport(self):
        if os.path.isfile(export_file):
            if sys.platform == 'win32':
                os.startfile(export_file)
            else:
                subprocess.call(['xdg-open', export_file])
        else:
            self.status_label.setText('Status: No export file, Export file first')


    def GetSetting(self):
        self.root_search_list = []
        self.search_exception_list = []
        f = open(search_locations, 'r')
        for line in f:
            self.root_search_list.append(line.rstrip('\n'))
        f.close()
        f = open(exception_file, 'r')
        for line in f:
            self.search_exception_list.append(line.rstrip('\n'))
        f.close()

    def SubCrawler(self, dir):
        QtGui.QApplication.processEvents()
        try:
            for filename in os.listdir(dir):
                path = os.path.join(dir, filename)
                if path not in self.search_exception_list:
                    # if the path is not a directory, then add it to the list
                    if not os.path.isdir(path):
                        self.pathlist.append(path)
                        self.file_count += 1
                        self.file_lcd.display(self.file_count)
                    else:
                        # if it is a directory, we must dive deeper
                        self.pathlist.append(path)
                        self.dir_count += 1
                        self.dir_lcd.display(self.dir_count)
                        self.SubCrawler(path)
        except:
            print('encountered an error with dir: ' + dir)
            if self.is_log_open:
                self.info.textedit.append('encountered an error with directory: ' + path)
            self.restricted_dir.append(dir)
            self.restricted_dir_count += 1


    def MainCrawler(self):
        start = time.time()
        for item in self.root_search_list:
            self.SubCrawler(item)
        self.pathlist.append(QtCore.QDateTime().currentDateTime().toString('MM.dd.yyyy - hh:mm:ss AP'))
        self.index_date.setText('Index -> Last updated: ' + self.pathlist[-1])
        end = time.time()
        print(end-start)

    def openText(self):
        self.editor = TextEditor()
        if self.deep_line != '':
            for item in self.listwidget.selectedItems():
                f = open(item.text(),'r')
                for line in f:
                    self.editor.textedit.append(line.rstrip('\n'))
                f.close()
                self.editor.highlight(self.deep_line.text())
                self.editor.show()

    def grep(self):
        text, ok = QtGui.QInputDialog.getText(self,'Grep','Enter Search Pattern')
        if ok:
            for item in self.listwidget.selectedItems():
                f = open(item.text(),'r')
                string = []
                for line in f:
                    string.append(line.rstrip('\n'))
                x=0
                xlist = []
                for item in string:
                    if text in item:
                        xlist.append(x)
                    x+=1

                self.editor = TextEditor()
                for x in xlist:
                    if x>1:
                        self.editor.textedit.append(string[x-1])
                    self.editor.textedit.append(string[x])
                    if x<len(string)-1:    
                        self.editor.textedit.append(string[x+1])
                    self.editor.textedit.append('\n')
                    x+=1

                self.editor.highlight(text)
                self.editor.show()

    def Search(self, term):
        print(term)
        self.listwidget.clear()
        self.found = []
        for path in self.pathlist:
            if term.lower() in path.lower():
                self.found.append(path)

    def SearchSplit(self):
        self.listwidget.clear()
        self.found = []
        if len(self.line_edit.text().split('|'))>1:
            term= self.line_edit.text().split('|')[0]
            text = self.line_edit.text()
            text = text[text.find('|')+1:]
            term2 = text.split('|')
            self.MultiSearch(term)
            for keys in term2:
                if keys!='':
                    self.SecondarySearch(keys)
        elif self.line_edit.text() == "@music":
            for path in self.pathlist:
                ext = os.path.splitext(path)[-1].lower()
                if ext in MUSIC_EXT:
                    self.found.append(path)

        elif self.line_edit.text() == "@video":
            for path in self.pathlist:
                ext = os.path.splitext(path)[-1].lower()
                if ext in VIDEO_EXT:
                    self.found.append(path)
        elif self.line_edit.text() == "@images":
            for path in self.pathlist:
                ext = os.path.splitext(path)[-1].lower()
                if ext in IMAGE_EXT:
                    self.found.append(path)
        elif self.line_edit.text() == '':
            self.listwidget.clear()
            self.found = []
            self.status_label.setText('Status: Cleared')  
        else:
            term = self.line_edit.text()
            self.MultiSearch(term)


        for item in self.found:
            self.listwidget.addItem(item)


        if len(self.found) == 0 and self.line_edit.text() != '':
            self.status_label.setText('Status: None Found...')
        if len(self.found) > 0:
            self.status_label.setText('Status: Found ' + str(len(self.found)) + ' files')



    def MultiSearch(self,term):
        if len(term.split(' ; ')) > 1 or len(term.split(' , ')) > 1:
            self.status_label.setText('Status: Please have no spaces for delimiters')
        elif len(term.split(';')) > 1 and '@' not in term:  # and filter
            filt = term.split(';')
            x = 1
            for fil in filt:
                templist = []
                if fil != '':
                    if x == 1:
                        for path in self.pathlist:
                            if fil.lower() in path.lower():
                                self.found.append(path)
                        x += 1
                    else:
                        for path in self.found:
                            if fil.lower() in path.lower():
                                templist.append(path)
                        self.found = []
                        for item in templist:
                            self.found.append(item)

        elif len(term.split(',')) > 1 and '@' not in term:  # or filter
            filt = term.split(',')

            for path in self.pathlist:
                if self.is_log_open:
                    self.info.textedit.append(path)
                for fil in filt:
                    if fil.lower() in path.lower():
                        if path not in self.found:
                            self.found.append(path)

        elif len(self.line_edit.text().split('@'))>1:
            self.hit = []
            type_ = self.line_edit.text().split("@")[1]
            search_key = self.line_edit.text().split("@")[0]
            if len(search_key.split(','))>1:
                filt = search_key.split(',')
                print(filt)
                for path in self.pathlist:
                    if self.is_log_open:
                        self.info.textedit.append(path)
                    for fil in filt:
                        if fil.lower() in path.lower():
                            if path not in self.hit:
                                self.hit.append(path)
            elif len(search_key.split(';'))>1:
                filt = search_key.split(';')
                print(filt)
                x = 1
                for fil in filt:
                    templist = []
                    if fil != '':
                        if x == 1:
                            for path in self.pathlist:
                                if fil.lower() in path.lower():
                                    self.hit.append(path)
                            x += 1
                        else:
                            for path in self.hit:
                                if fil.lower() in path.lower():
                                    templist.append(path)
                            self.hit = []
                            for item in templist:
                                self.hit.append(item)
            else:
                for path in self.pathlist:
                    if search_key.lower() in path.lower():
                        self.hit.append(path)
            if type_ =='music':
                for item in self.hit:
                    if os.path.splitext(item)[-1].lower() in MUSIC_EXT:
                        self.found.append(item)
            elif type_ == 'video':
                for item in self.hit:
                    if os.path.splitext(item)[-1].lower() in VIDEO_EXT:
                        self.found.append(item)
            elif type_ == 'images':
                for item in self.hit:
                    if os.path.splitext(item)[-1].lower() in IMAGE_EXT:
                        self.found.append(item)


        else:
            filt = term.rstrip('|')
            for path in self.pathlist:
                if filt.lower() in path.lower():
                    self.found.append(path)

    def SecondarySearch(self,term):
        templist = []
        if len(term.split(' ; ')) > 1 or len(term.split(' , ')) > 1:
            self.status_label.setText('Status: Please have no spaces for delimiters')
        elif len(term.split(';')) > 1:  # and filter
            filt = term.split(';')
            x = 1
            for fil in filt:
                templist2 = []
                if fil != '':
                    if x == 1:
                        for path in self.found:
                            if fil.lower() in path.lower():
                                templist.append(path)
                        x += 1
                    else:
                        for path in templist:
                            if fil.lower() in path.lower():
                                templist2.append(path)
                        templist = []
                        for item in templist2:
                            templist.append(item)
        elif len(term.split(',')) > 1:  # or filter
            filt = term.split(',')
            print(filt)
            for path in self.found:
                for fil in filt:
                    if fil.lower() in path.lower():
                        if path not in templist:
                            templist.append(path)
        else:
            for path in self.found:
                if term.lower() in path.lower():
                    templist.append(path)

        self.found = []
        for item in templist:
            self.found.append(item)

    # still experimenting: only going to do text files
    def DeepSearch(self):
        self.listwidget.clear()
        starttime = time.time()
        if self.deep_line.text() == '':
            self.listwidget.clear()
            self.status_label.setText('Status: Cleared')
        elif self.line_edit.text()!='' and len(self.found)>0:
            total = len(self.found)
            term = self.deep_line.text()
            self.checkTextFromResult(term,total)
        elif self.line_edit.text()!='' and len(self.found)==0:
            self.SearchSplit()
            total = len(self.found)
            term = self.deep_line.text()
            self.checkTextFromResult(term,total)
        else:
            total = len(self.pathlist)
            count = 0
            self.status_label.setText('Status: Searching...')
            for path in self.pathlist:
                if count%50 == 0:
                    QtGui.QApplication.processEvents()
                count += 1
                self.progress_bar.setValue(int(100 * (count / total)))
                if self.is_log_open:
                    self.info.textedit.append(path)
                # self.status_label.setText('Status: processing -> ' + path)
                ext = os.path.splitext(path)[-1].lower()
                if ext in EXT_LIST:
                    try:
                        if os.path.isfile(path):
                            f = open(path, 'r')
                            text = f.read()
                            if self.deep_line.text().lower() in text.lower():
                                self.found.append(path)
                                self.listwidget.addItem(path)
                            f.close()
                    except (UnicodeDecodeError, PermissionError) as e:
                        if self.is_log_open:
                            self.info.textedit.append('error: ' + path)
                        self.error_list.append('error: ' + path)
                if ext in ['.docx']:
                    if self.is_log_open:
                        self.info.textedit.append('Checking: ' + path)
                    try:
                        document = docx.Document(path)
                        templist = []
                        for paragraph in document.paragraphs:
                            if self.deep_line.text().lower() in paragraph.text.lower():
                                if path not in templist:
                                    self.listwidget.addItem(path)
                                    templist.append(path)
                    except (ValueError, docx.opc.exceptions.PackageNotFoundError) as e:
                        print(repr(e))
                        print('error')

        if self.listwidget.count() == 0:
            self.status_label.setText('Status: None Found...')
        else:
            self.status_label.setText('Status: Found ' + str(self.listwidget.count()) + ' files')
        endtime = time.time()
        print(endtime - starttime)


    def checkTextFromResult(self,term,total):
        self.listwidget.clear()
        count = 0
        for path in self.found:
                QtGui.QApplication.processEvents()
                count += 1
                self.progress_bar.setValue(int(100 * (count / total)))
                if self.is_log_open:
                    self.info.textedit.append(path)
                self.status_label.setText('Status: processing ->' + path)
                ext = os.path.splitext(path)[-1].lower()
                if ext in EXT_LIST:
                    if self.is_log_open:
                        self.info.textedit.append('Checking: ' + path)
                    try:
                        if os.path.isfile(path):
                            f = open(path, 'r')
                            text = f.read()
                            if term.lower() in text.lower():
                                self.listwidget.addItem(path)
                            f.close()
                    except UnicodeDecodeError:
                        print('error: ' + path)
                        if self.is_log_open:
                            self.info.textedit.append('error: ' + path)
                        self.error_list.append('error: ' + path)
                if ext in ['.docx']:
                    if self.is_log_open:
                        self.info.textedit.append('Checking: ' + path)
                    try:
                        document = docx.Document(path)
                        templist = []
                        for paragraph in document.paragraphs:
                            if term.lower() in paragraph.text.lower():
                                if path not in templist:
                                    self.listwidget.addItem(path)
                                    templist.append(path)
                    except ValueError:
                        print(ValueError)

    def OpenFile(self):
        for item in self.listwidget.selectedItems():
            if sys.platform == 'win32':
                os.startfile(item.text())
            else:
                app_type = self.xdgQuery('filetype', item.text())
                default = self.xdgQuery('default', app_type)
                print(app_type, default)
                if default == b'':
                    print('default is none')
                    app = self.inputDialog(app_type)
                    print(app)
                    self.setXdgDefault(app, app_type)
                default = self.xdgQuery('default', app_type)
                if default != b'':
                    self.thread = ProcessThread(item.text())
                    self.thread.start()

    def xdgQuery(self, command, path):
        p = subprocess.Popen(['xdg-mime', 'query', command, path],
                             stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        output, errors = p.communicate()
        if p.returncode or errors:
            raise XDGError('xdg-mime returned error code %d: %s' %
                           (p.returncode, errors.strip()))
        if command == 'default':
            return output
        return output.strip().decode('utf-8')

    def setXdgDefault(self, app, app_type):
        subprocess.call(['xdg-mime', 'default', app, app_type])

    def inputDialog(self, app_type):
        dia = 'no/invalid default app for: ' + app_type + '-> set one below (ex:gedit)\n (refer to view->app table if you are not sure)'
        text, ok = QtGui.QInputDialog.getText(self, 'set default application', dia)
        if ok:
            if self.checkApps(text):
                return text + '.desktop'
            else:
                self.inputDialog(app_type)

    def checkApps(self, text):
        for item in os.listdir(r'/usr/share/applications'):
            if text == item.split('.')[0]:
                return True
        return False


    def OpenRestricted(self):
        if sys.platform == 'win32':
            os.startfile(restricted_file)
        else:
            subprocess.call(['xdg-open', restricted_file])


    def ReadArchive(self):
        f = open(path_archive, 'r')
        for line in f:
            self.pathlist.append(line.rstrip('\n'))
        self.status_label.setText('Status: Index imported, ready for use!')

    def GenerateArchive(self):
        self.status_label.setText('Status: Generating Path Index...')
        if self.is_log_open:
            self.info.textedit.clear()
        self.file_count = 0
        self.dir_count = 0
        self.pathlist = []
        self.restricted_dir = []
        self.error_list = []
        if os.path.isfile(path_archive):
            os.remove(path_archive)
        self.GetSetting()
        self.MainCrawler()
        if self.is_log_open:
            self.info.textedit.append('\n')
            self.info.textedit.append('Errors: ' + str(len(self.error_list)) + '\n')
            for item in self.error_list:
                self.info.textedit.append('| ' + item)
            self.info.textedit.append('\n')
            self.info.textedit.append('Directories not accessed: ' + str(self.restricted_dir_count) + '\n')
            for item in self.restricted_dir:
                self.info.textedit.append('|| ' + item)
        self.CreateArchive()
        self.CreateRestrictedFolder()
        self.dir_lcd.display(self.dir_count)
        self.file_lcd.display(self.file_count)
        self.status_label.setText('Status: Completed, Index archived and ready for use!')

    def CreateArchive(self):
        date_time = QtCore.QDateTime().currentDateTime().toString('MM.dd.yyyy - hh:mm:ss AP')
        f = open(path_archive, 'w+')
        x = 0
        for a in self.pathlist:
            try:
                if x ==0:
                    f.write(date_time)
                    f.write('\n')
                    x+=1
                f.write(a)
                f.write("\n")
            except UnicodeEncodeError:
                error = 'couldnt add: ' + a + ' to file'
                self.error_list.append(error)
                if self.is_log_open:
                    self.info.textedit.append(error)
        f.close()
        f = open(file_count, 'w+')
        for item in [self.dir_count, self.file_count]:
            f.write(str(item))
            f.write('\n')
        f.close()

    def CreateRestrictedFolder(self):
        f = open(restricted_file, 'w+')
        for item in self.restricted_dir:
            try:
                f.write(item)
                f.write('\n')
            except UnicodeEncodeError:
                print(UnicodeEncodeError)
        f.close()

    # def createWordList(self):
    #     self.wordlist = []
    #     for item in self.pathlist:
    #         QtGui.QApplication.processEvents()
    #         if sys.platform == 'win32':
    #             terms = item.split('\\')
    #         else:
    #             terms = item.split('/')
    #         for word in terms:
    #             if word not in self.wordlist:
    #                 print('appending: ' + word)
    #                 self.wordlist.append(word)
    #     f = open(word_list_file,'w+')
    #     for item in self.wordlist:
    #         try:
    #             f.write(item)
    #             f.write('\n')
    #         except:
    #             print('error writing word to file.')
    #     f.close()


class XDGError(Exception): pass


def main():
    app = QtGui.QApplication(sys.argv)
    search = Searcher()
    sys.exit(app.exec_())


if __name__ == '__main__':
    main()
